"""Word 全功能展示文档 - python-docx 能力一览"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ====== 配色方案：奶酪色 × 蒂芙尼蓝 ======
TIFFANY = RGBColor(0x81, 0xD8, 0xD1)   # 蒂芙尼蓝 #81D8D1
DEEP_TF = RGBColor(0x5C, 0xBD, 0xB4)   # 深蒂蓝 #5CBDB4
CREAM    = RGBColor(0xFB, 0xF7, 0xD8)   # 奶酪色 #FBF7D8
LIGHT_CR = RGBColor(0xFD, 0xF9, 0xEC)   # 浅奶黄 #FDF9EC
MINT     = RGBColor(0xA8, 0xE0, 0xD8)   # 薄荷绿 #A8E0D8
DARK     = RGBColor(0x2B, 0x2D, 0x42)   # 正文深色
GRAY     = RGBColor(0x7A, 0x7C, 0x8E)   # 灰色次要文字
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


# ========== 辅助函数 ==========

def set_cell_bg(cell, color_hex):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """设置单元格边框
    kwargs: top, bottom, left, right - dict with sz, color, val
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} '
            f'w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" '
            f'w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_paragraph_spacing(para, before=0, after=0, line=None):
    """设置段落间距"""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)


def add_horizontal_line(doc, color="81D8D1"):
    """添加水平分割线"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_colored_heading(doc, text, level=1, color=None):
    """添加彩色标题"""
    if color is None:
        color = TIFFANY
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color
    return heading


def add_styled_paragraph(doc, text, font_size=12, bold=False, italic=False,
                         color=None, align=None, indent=None, line_spacing=None):
    """添加样式段落"""
    if color is None:
        color = DARK
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Pt(indent)
    if line_spacing:
        p.paragraph_format.line_spacing = Pt(line_spacing)
    return p


# ========== 文档生成 ==========

doc = Document()

# --- 页面设置 ---
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

# --- 默认字体 ---
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(12)
font.color.rgb = DARK
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


# ========================================
# 第1部分：封面
# ========================================
for _ in range(6):
    doc.add_paragraph()

# 封面装饰线
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="24" w:space="1" w:color="81D8D1"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

# 主标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Python-docx 全功能展示")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = DEEP_TF
run.font.name = "Microsoft YaHei"
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# 副标题
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Word 文档生成能力一览")
run.font.size = Pt(18)
run.font.color.rgb = GRAY
run.font.name = "Microsoft YaHei"
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# 封面装饰线
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:bottom w:val="single" w:sz="24" w:space="1" w:color="81D8D1"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

for _ in range(4):
    doc.add_paragraph()

# 作者信息
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Hinata | 2026.06")
run.font.size = Pt(14)
run.font.color.rgb = GRAY

# 分页
doc.add_page_break()


# ========================================
# 第2部分：目录占位
# ========================================
add_colored_heading(doc, "目录", level=1, color=DEEP_TF)
add_horizontal_line(doc)

# 插入 TOC 域代码（需要在 Word 中按 F9 刷新）
p = doc.add_paragraph()
run = p.add_run()
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._r.append(fldChar1)
run2 = p.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
run2._r.append(instrText)
run3 = p.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
run3._r.append(fldChar2)
run4 = p.add_run("（请在 Word 中右键 → 更新域 → 更新整个目录）")
run4.font.color.rgb = GRAY
run4.font.size = Pt(10)
run5 = p.add_run()
fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run5._r.append(fldChar3)

doc.add_page_break()


# ========================================
# 第3部分：文字排版
# ========================================
add_colored_heading(doc, "一、文字排版", level=1, color=DEEP_TF)
add_horizontal_line(doc)

# --- 字号对比 ---
add_colored_heading(doc, "1.1 字号对比", level=2, color=TIFFANY)

sizes = [
    ("36pt - 大标题", 36, True),
    ("24pt - 章节标题", 24, True),
    ("18pt - 小节标题", 18, True),
    ("14pt - 副标题", 14, True),
    ("12pt - 正文（默认）", 12, False),
    ("10pt - 小字注释", 10, False),
    ("8pt - 脚注", 8, False),
]

for text, size, bold in sizes:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = DARK
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    set_paragraph_spacing(p, before=4, after=4)

doc.add_paragraph()

# --- 颜色展示 ---
add_colored_heading(doc, "1.2 颜色展示", level=2, color=TIFFANY)

colors = [
    ("蒂芙尼蓝 - 经典 Tiffany 蓝", TIFFANY),
    ("深蒂蓝 - 深色装饰", DEEP_TF),
    ("奶酪色文字 - 温暖奶黄", RGBColor(0xC0, 0xB0, 0x60)),
    ("薄荷绿 - 中间色", MINT),
    ("灰色 - 次要信息", GRAY),
]

for text, color in colors:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = color
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    set_paragraph_spacing(p, before=4, after=4)

doc.add_paragraph()

# --- 字体样式 ---
add_colored_heading(doc, "1.3 字体样式", level=2, color=TIFFANY)

styles = [
    ("加粗文本", True, False, False),
    ("斜体文本", False, True, False),
    ("粗斜体文本", True, True, False),
    ("下划线文本", False, False, "underline"),
    ("删除线文本", False, False, "strike"),
]

for text, bold, italic, extra in styles:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = DARK
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if extra == "underline":
        run.font.underline = True
    elif extra == "strike":
        run.font.strike = True
    set_paragraph_spacing(p, before=4, after=4)

doc.add_paragraph()

# --- 高亮 ---
add_colored_heading(doc, "1.4 高亮效果", level=2, color=TIFFANY)

p = doc.add_paragraph()
run = p.add_run("这段文字有")
run.font.size = Pt(14)
run.font.color.rgb = DARK
run.font.name = "Microsoft YaHei"
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

run2 = p.add_run("蒂芙尼蓝高亮")
run2.font.size = Pt(14)
run2.font.color.rgb = DARK
run2.font.name = "Microsoft YaHei"
run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="81D8D1"/>')
run2._element.rPr.append(shd)

run3 = p.add_run("效果")
run3.font.size = Pt(14)
run3.font.color.rgb = DARK
run3.font.name = "Microsoft YaHei"
run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_page_break()


# ========================================
# 第4部分：段落排版
# ========================================
add_colored_heading(doc, "二、段落排版", level=1, color=DEEP_TF)
add_horizontal_line(doc)

# --- 对齐方式 ---
add_colored_heading(doc, "2.1 对齐方式", level=2, color=TIFFANY)

aligns = [
    ("左对齐（默认）", WD_ALIGN_PARAGRAPH.LEFT),
    ("居中对齐", WD_ALIGN_PARAGRAPH.CENTER),
    ("右对齐", WD_ALIGN_PARAGRAPH.RIGHT),
    ("两端对齐", WD_ALIGN_PARAGRAPH.JUSTIFY),
]

for text, align in aligns:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = DARK
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p.alignment = align
    set_paragraph_spacing(p, before=4, after=4)

doc.add_paragraph()

# --- 行距 ---
add_colored_heading(doc, "2.2 行距控制", level=2, color=TIFFANY)

line_spacings = [
    ("单倍行距（默认）", 15),
    ("1.5 倍行距", 22),
    ("双倍行距", 30),
    ("固定值 40pt", 40),
]

text_content = "这是示例段落，用于展示不同行距效果。Word 文档的排版能力远超普通文本编辑器。"

for label, spacing in line_spacings:
    p = doc.add_paragraph()
    run = p.add_run(f"【{label}】")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = TIFFANY
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run2 = p.add_run(text_content)
    run2.font.size = Pt(12)
    run2.font.color.rgb = DARK
    run2.font.name = "Microsoft YaHei"
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    set_paragraph_spacing(p, before=6, after=6, line=spacing)

doc.add_paragraph()

# --- 缩进 ---
add_colored_heading(doc, "2.3 首行缩进", level=2, color=TIFFANY)

for indent, label in [(0, "无缩进"), (24, "首行缩进 2 字符（中文标准）"), (48, "首行缩进 4 字符")]:
    p = doc.add_paragraph()
    run = p.add_run(f"【{label}】")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = TIFFANY
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run2 = p.add_run("这是一段较长的示例文字，用于展示首行缩进的效果。中文排版通常使用首行缩进两字符的格式。这段文字会自动换行，可以看到缩进效果。")
    run2.font.size = Pt(12)
    run2.font.color.rgb = DARK
    run2.font.name = "Microsoft YaHei"
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p.paragraph_format.first_line_indent = Pt(indent)
    set_paragraph_spacing(p, before=4, after=4)

doc.add_page_break()


# ========================================
# 第5部分：列表
# ========================================
add_colored_heading(doc, "三、列表", level=1, color=DEEP_TF)
add_horizontal_line(doc)

# --- 无序列表 ---
add_colored_heading(doc, "3.1 无序列表", level=2, color=TIFFANY)

bullets = [
    "python-docx 可以生成 Word 文档",
    "支持 .docx 格式",
    "纯 Python 实现，无需安装 Word",
    "开源库，社区活跃",
]

for item in bullets:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = DARK
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph()

# --- 有序列表 ---
add_colored_heading(doc, "3.2 有序列表", level=2, color=TIFFANY)

steps = [
    "安装 python-docx：pip install python-docx",
    "创建 Document 对象",
    "添加段落、表格、图片等内容",
    "保存为 .docx 文件",
]

for item in steps:
    p = doc.add_paragraph(item, style='List Number')
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = DARK
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_page_break()


# ========================================
# 第6部分：表格
# ========================================
add_colored_heading(doc, "四、表格", level=1, color=DEEP_TF)
add_horizontal_line(doc)

# --- 基础表格 ---
add_colored_heading(doc, "4.1 基础表格", level=2, color=TIFFANY)

data1 = [
    ["姓名", "职位", "部门", "入职年份"],
    ["张三", "高级工程师", "技术部", "2020"],
    ["李四", "产品经理", "产品部", "2021"],
    ["王五", "UI 设计师", "设计部", "2022"],
    ["赵六", "数据分析师", "数据部", "2023"],
]

table1 = doc.add_table(rows=len(data1), cols=4, style='Table Grid')
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(data1):
    row = table1.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.font.size = Pt(11)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        if i == 0:
            # 表头
            run.font.bold = True
            run.font.color.rgb = WHITE
            set_cell_bg(cell, "81D8D1")
        else:
            run.font.color.rgb = DARK
            if i % 2 == 0:
                set_cell_bg(cell, "FDF9EC")

doc.add_paragraph()

# --- 合并单元格 ---
add_colored_heading(doc, "4.2 合并单元格", level=2, color=TIFFANY)

table2 = doc.add_table(rows=4, cols=4, style='Table Grid')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

# 表头合并
a = table2.cell(0, 0)
b = table2.cell(0, 3)
a.merge(b)
p = a.paragraphs[0]
p.text = ""
run = p.add_run("项目汇总表")
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = WHITE
run.font.name = "Microsoft YaHei"
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
set_cell_bg(a, "5CBDB4")

# 数据行（表头行已在 row 0）
data2 = [
    ["前端开发", "张三", "120h", "进行中"],
    ["后端开发", "李四", "200h", "已完成"],
    ["测试验收", "王五", "80h", "待开始"],
]

for i, row_data in enumerate(data2):
    row = table2.rows[i + 1]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.font.size = Pt(11)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run.font.color.rgb = DARK

        if i % 2 == 0:
            set_cell_bg(cell, "FDF9EC")

doc.add_paragraph()

# --- 带底色表格 ---
add_colored_heading(doc, "4.3 学习路线表（交替色）", level=2, color=TIFFANY)

roadmap = [
    ["阶段", "内容", "工具", "产出"],
    ["第1周", "Python 基础", "VS Code", "爬虫小 demo"],
    ["第2周", "Web 框架", "Flask/Django", "个人博客"],
    ["第3周", "数据库", "MySQL/Redis", "数据存储方案"],
    ["第4周", "部署上线", "Docker/Nginx", "线上可访问"],
]

table3 = doc.add_table(rows=len(roadmap), cols=4, style='Table Grid')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(roadmap):
    row = table3.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.font.size = Pt(11)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        if i == 0:
            run.font.bold = True
            run.font.color.rgb = WHITE
            set_cell_bg(cell, "81D8D1")
        else:
            run.font.color.rgb = DARK
            if i % 2 == 1:
                set_cell_bg(cell, "FDF9EC")

doc.add_page_break()


# ========================================
# 第7部分：页眉页脚
# ========================================
add_colored_heading(doc, "五、页眉页脚", level=1, color=DEEP_TF)
add_horizontal_line(doc)

add_styled_paragraph(doc, "本节展示页眉页脚和页码效果。", font_size=12, color=DARK)
add_styled_paragraph(doc, "请在 Word 中查看页面顶部和底部。", font_size=12, color=GRAY)

# --- 页眉 ---
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.text = ""
run = hp.add_run("Python-docx 全功能展示")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.font.name = "Microsoft YaHei"
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# 页眉底部线
pPr = hp._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="81D8D1"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

# --- 页脚 + 页码 ---
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.text = ""

# 页码域代码
run1 = fp.add_run("— ")
run1.font.size = Pt(9)
run1.font.color.rgb = GRAY

run2 = fp.add_run()
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run2._r.append(fldChar1)
run3 = fp.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
run3._r.append(instrText)
run4 = fp.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
run4._r.append(fldChar2)
run5 = fp.add_run("1")
run5.font.size = Pt(9)
run5.font.color.rgb = GRAY
run6 = fp.add_run()
fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run6._r.append(fldChar3)

run7 = fp.add_run(" —")
run7.font.size = Pt(9)
run7.font.color.rgb = GRAY

fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()


# ========================================
# 第8部分：分栏
# ========================================
add_colored_heading(doc, "六、分节与分栏", level=1, color=DEEP_TF)
add_horizontal_line(doc)

add_styled_paragraph(doc, "Word 支持在同一文档中使用不同的页面布局。", font_size=12, color=DARK)
add_styled_paragraph(doc, "本页为纵向（Portrait），下一页将切换为横向（Landscape）。", font_size=12, color=GRAY)

doc.add_page_break()

# --- 横向页面 ---
new_section = doc.add_section()
new_section.orientation = WD_ORIENT.LANDSCAPE
new_section.page_width = Cm(29.7)
new_section.page_height = Cm(21)

add_colored_heading(doc, "横向页面", level=1, color=DEEP_TF)
add_horizontal_line(doc)

# 横向页面放一个宽表格
data_wide = [
    ["功能", "python-docx", "win32com", "难度", "推荐度"],
    ["段落排版", "✅ 完美", "✅ 可用", "⭐", "⭐⭐⭐⭐⭐"],
    ["表格创建", "✅ 完美", "✅ 可用", "⭐", "⭐⭐⭐⭐⭐"],
    ["页眉页脚", "✅ 完美", "✅ 可用", "⭐⭐", "⭐⭐⭐⭐⭐"],
    ["目录刷新", "❌ 不支持", "✅ 可用", "⭐⭐", "⭐⭐⭐⭐"],
    ["水印", "❌ 不支持", "✅ 可用", "⭐⭐⭐", "⭐⭐⭐"],
    ["批注修订", "❌ 不支持", "✅ 可用", "⭐⭐⭐", "⭐⭐⭐"],
    ["邮件合并", "❌ 不支持", "✅ 可用", "⭐⭐⭐", "⭐⭐"],
]

table_wide = doc.add_table(rows=len(data_wide), cols=5, style='Table Grid')
table_wide.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(data_wide):
    row = table_wide.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.font.size = Pt(11)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        if i == 0:
            run.font.bold = True
            run.font.color.rgb = WHITE
            set_cell_bg(cell, "81D8D1")
        else:
            run.font.color.rgb = DARK
            if i % 2 == 0:
                set_cell_bg(cell, "FDF9EC")

# 设置列宽
for row in table_wide.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(4)
    row.cells[2].width = Cm(4)
    row.cells[3].width = Cm(3)
    row.cells[4].width = Cm(3)

# 恢复纵向
new_section2 = doc.add_section()
new_section2.orientation = WD_ORIENT.PORTRAIT
new_section2.page_width = Cm(21)
new_section2.page_height = Cm(29.7)

add_colored_heading(doc, "回到纵向页面", level=2, color=TIFFANY)
add_styled_paragraph(doc, "文档可以自由切换纵向和横向，适用于报告、方案等需要大表格的场景。", font_size=12, color=DARK)

doc.add_page_break()


# ========================================
# 第9部分：分隔线
# ========================================
add_colored_heading(doc, "七、分隔线与装饰", level=1, color=DEEP_TF)
add_horizontal_line(doc)

add_colored_heading(doc, "蒂芙尼蓝分隔线", level=2, color=TIFFANY)
add_horizontal_line(doc, "81D8D1")

add_colored_heading(doc, "深蒂蓝分隔线", level=2, color=TIFFANY)
add_horizontal_line(doc, "5CBDB4")

add_colored_heading(doc, "薄荷绿分隔线", level=2, color=TIFFANY)
add_horizontal_line(doc, "A8E0D8")

add_colored_heading(doc, "灰色分隔线", level=2, color=TIFFANY)
add_horizontal_line(doc, "7A7C8E")

doc.add_paragraph()
add_styled_paragraph(doc, "分隔线可以用于区分文档的不同章节，增强可读性。", font_size=12, color=DARK)

doc.add_page_break()


# ========================================
# 第10部分：混排
# ========================================
add_colored_heading(doc, "八、混排示例", level=1, color=DEEP_TF)
add_horizontal_line(doc)

add_styled_paragraph(doc, "以下展示实际文档中常见的混排效果：标题 + 正文 + 表格 + 列表。", font_size=12, color=GRAY)

doc.add_paragraph()

# 小节1
add_colored_heading(doc, "项目背景", level=2, color=TIFFANY)

text1 = "本项目旨在开发一个基于大语言模型的智能客服系统，支持多轮对话、意图识别和自动回复。"
p = doc.add_paragraph()
run = p.add_run(text1)
run.font.size = Pt(12)
run.font.color.rgb = DARK
run.font.name = "Microsoft YaHei"
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
p.paragraph_format.first_line_indent = Pt(24)
set_paragraph_spacing(p, before=6, after=6)

text2 = "项目周期为 3 个月，预算 50 万元，团队规模 8 人。预计 Q3 上线内测版本。"
p = doc.add_paragraph()
run = p.add_run(text2)
run.font.size = Pt(12)
run.font.color.rgb = DARK
run.font.name = "Microsoft YaHei"
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
p.paragraph_format.first_line_indent = Pt(24)
set_paragraph_spacing(p, before=6, after=6)

doc.add_paragraph()

# 小节2
add_colored_heading(doc, "团队分工", level=2, color=TIFFANY)

team = [
    ["角色", "人数", "职责"],
    ["项目经理", "1", "统筹协调、进度管理"],
    ["后端开发", "3", "API 开发、模型集成"],
    ["前端开发", "2", "界面开发、交互设计"],
    ["测试工程师", "1", "功能测试、性能测试"],
    ["运维工程师", "1", "部署、监控、运维"],
]

table_team = doc.add_table(rows=len(team), cols=3, style='Table Grid')
table_team.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, row_data in enumerate(team):
    row = table_team.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.font.size = Pt(11)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        if i == 0:
            run.font.bold = True
            run.font.color.rgb = WHITE
            set_cell_bg(cell, "81D8D1")
        else:
            run.font.color.rgb = DARK
            if i % 2 == 0:
                set_cell_bg(cell, "FDF9EC")

doc.add_paragraph()

# 小节3
add_colored_heading(doc, "关键里程碑", level=2, color=TIFFANY)

milestones = [
    "6月15日 - 需求评审通过",
    "6月30日 - 技术方案确定",
    "7月15日 - 核心功能开发完成",
    "7月31日 - 内部测试",
    "8月15日 - 公测上线",
]

for item in milestones:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = DARK
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


# ========================================
# 保存
# ========================================
output_path = os.path.join(OUTPUT_DIR, "python_docx_全功能展示.docx")
doc.save(output_path)
print(f"Word 文档已生成: {output_path}")
print("python-docx 全功能展示完成!")
